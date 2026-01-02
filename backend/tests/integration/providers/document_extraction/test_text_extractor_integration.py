import pytest
from docx import Document
import io

from packages.documents.providers.document_extraction.text_extractor import (
    TextExtractor,
)


class TestTextExtractorIntegration:
    """Integration tests for TextExtractor that use real document processing."""

    @pytest.fixture
    def text_extractor(self):
        """Create TextExtractor instance."""
        return TextExtractor()

    @pytest.fixture
    def sample_pdf_content(self):
        """Sample PDF content for testing (minimal PDF)."""
        # This is a minimal valid PDF content
        return b"""%PDF-1.4
1 0 obj
<<
/Type /Catalog
/Pages 2 0 R
>>
endobj

2 0 obj
<<
/Type /Pages
/Kids [3 0 R]
/Count 1
>>
endobj

3 0 obj
<<
/Type /Page
/Parent 2 0 R
/Resources <<
/Font <<
/F1 4 0 R
>>
>>
/MediaBox [0 0 612 792]
/Contents 5 0 R
>>
endobj

4 0 obj
<<
/Type /Font
/Subtype /Type1
/BaseFont /Times-Roman
>>
endobj

5 0 obj
<<
/Length 44
>>
stream
BT
/F1 12 Tf
72 720 Td
(Hello World) Tj
ET
endstream
endobj

xref
0 6
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000274 00000 n 
0000000354 00000 n 
trailer
<<
/Size 6
/Root 1 0 R
>>
startxref
451
%%EOF"""

    @pytest.fixture
    def sample_docx_content(self):
        """Sample DOCX content for testing."""
        # Create a minimal DOCX file in memory using python-docx

        doc = Document()
        doc.add_paragraph("This is a test document.")
        doc.add_paragraph("It contains multiple paragraphs.")
        doc.add_paragraph("For testing text extraction.")

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    @pytest.fixture
    def sample_text_content(self):
        """Sample plain text content."""
        return b"This is a plain text document.\nIt has multiple lines.\nFor testing purposes."

    @pytest.fixture
    def sample_markdown_content(self):
        """Sample markdown content."""
        return b"""# Test Document

This is a **markdown** document for testing.

## Features

- Item 1
- Item 2  
- Item 3

### Code Example

```python
def hello():
    print("Hello World")
```

End of document."""

    def test_supports_file_type(self, text_extractor):
        """Test file type support detection."""
        # Supported types
        assert text_extractor.supports_file_type("application/pdf")
        assert text_extractor.supports_file_type(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        assert text_extractor.supports_file_type("text/plain")
        assert text_extractor.supports_file_type("text/markdown")

        # Case insensitive
        assert text_extractor.supports_file_type("APPLICATION/PDF")
        assert text_extractor.supports_file_type("TEXT/PLAIN")

        # Unsupported types
        assert not text_extractor.supports_file_type("image/jpeg")
        assert not text_extractor.supports_file_type("application/zip")
        assert not text_extractor.supports_file_type("video/mp4")

    @pytest.mark.asyncio
    async def test_extract_pdf_text_integration(
        self, text_extractor, sample_pdf_content
    ):
        """Test PDF text extraction with real PDF processing."""
        text = await text_extractor.extract_text(
            file_data=sample_pdf_content, file_type="application/pdf"
        )

        assert isinstance(text, str)
        assert len(text) > 0
        assert "Hello World" in text
        print(f"Extracted PDF text: {text}")

    @pytest.mark.asyncio
    async def test_extract_docx_text_integration(
        self, text_extractor, sample_docx_content
    ):
        """Test DOCX text extraction with real DOCX processing."""
        text = await text_extractor.extract_text(
            file_data=sample_docx_content,
            file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        assert isinstance(text, str)
        assert len(text) > 0
        assert "test document" in text.lower()
        assert "multiple paragraphs" in text.lower()
        assert "text extraction" in text.lower()
        print(f"Extracted DOCX text: {text}")

    @pytest.mark.asyncio
    async def test_extract_plain_text_integration(
        self, text_extractor, sample_text_content
    ):
        """Test plain text extraction."""
        text = await text_extractor.extract_text(
            file_data=sample_text_content, file_type="text/plain"
        )

        assert isinstance(text, str)
        assert len(text) > 0
        assert "plain text document" in text
        assert "multiple lines" in text
        print(f"Extracted plain text: {text}")

    @pytest.mark.asyncio
    async def test_extract_markdown_text_integration(
        self, text_extractor, sample_markdown_content
    ):
        """Test markdown text extraction."""
        text = await text_extractor.extract_text(
            file_data=sample_markdown_content, file_type="text/markdown"
        )

        assert isinstance(text, str)
        assert len(text) > 0
        assert "Test Document" in text
        assert "markdown" in text
        assert "def hello():" in text
        print(f"Extracted markdown text: {text}")

    @pytest.mark.asyncio
    async def test_get_pdf_metadata_integration(
        self, text_extractor, sample_pdf_content
    ):
        """Test PDF metadata extraction with real PDF processing."""
        metadata = await text_extractor.get_metadata(
            file_data=sample_pdf_content, file_type="application/pdf"
        )

        assert isinstance(metadata, dict)
        assert "file_size" in metadata
        assert "file_type" in metadata
        assert "page_count" in metadata
        assert metadata["file_type"] == "application/pdf"
        assert metadata["page_count"] >= 1
        assert metadata["file_size"] == len(sample_pdf_content)
        print(f"PDF metadata: {metadata}")

    @pytest.mark.asyncio
    async def test_get_docx_metadata_integration(
        self, text_extractor, sample_docx_content
    ):
        """Test DOCX metadata extraction with real DOCX processing."""
        metadata = await text_extractor.get_metadata(
            file_data=sample_docx_content,
            file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        assert isinstance(metadata, dict)
        assert "file_size" in metadata
        assert "file_type" in metadata
        assert "paragraph_count" in metadata
        assert (
            metadata["file_type"]
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        assert metadata["paragraph_count"] >= 3  # We created 3 paragraphs
        assert metadata["file_size"] == len(sample_docx_content)
        print(f"DOCX metadata: {metadata}")

    @pytest.mark.asyncio
    async def test_get_plain_text_metadata_integration(
        self, text_extractor, sample_text_content
    ):
        """Test plain text metadata extraction."""
        metadata = await text_extractor.get_metadata(
            file_data=sample_text_content, file_type="text/plain"
        )

        assert isinstance(metadata, dict)
        assert "file_size" in metadata
        assert "file_type" in metadata
        assert metadata["file_type"] == "text/plain"
        assert metadata["file_size"] == len(sample_text_content)
        print(f"Plain text metadata: {metadata}")

    @pytest.mark.asyncio
    async def test_unsupported_file_type(self, text_extractor):
        """Test handling of unsupported file types."""
        with pytest.raises(Exception) as exc_info:
            await text_extractor.extract_text(
                file_data=b"fake image data", file_type="image/jpeg"
            )

        assert "Unsupported file type" in str(exc_info.value)

    @pytest.mark.asyncio
    async def test_corrupted_pdf_handling(self, text_extractor):
        """Test handling of corrupted PDF files."""
        corrupted_pdf = b"Not a real PDF file"

        with pytest.raises(Exception) as exc_info:
            await text_extractor.extract_text(
                file_data=corrupted_pdf, file_type="application/pdf"
            )

        assert "Failed to extract text" in str(exc_info.value)

    @pytest.mark.asyncio
    async def test_empty_file_handling(self, text_extractor):
        """Test handling of empty files."""
        empty_content = b""

        text = await text_extractor.extract_text(
            file_data=empty_content, file_type="text/plain"
        )

        assert isinstance(text, str)
        assert len(text) == 0

    @pytest.mark.asyncio
    async def test_unicode_text_handling(self, text_extractor):
        """Test handling of Unicode text content."""
        unicode_content = "Hello 世界 🌍 Héllo Wörld".encode("utf-8")

        text = await text_extractor.extract_text(
            file_data=unicode_content, file_type="text/plain"
        )

        assert isinstance(text, str)
        assert "世界" in text
        assert "🌍" in text
        assert "Héllo" in text
        assert "Wörld" in text
        print(f"Unicode text: {text}")

    @pytest.mark.asyncio
    async def test_latin1_fallback_handling(self, text_extractor):
        """Test Latin-1 fallback for text that fails UTF-8 decoding."""
        # Create content that's valid Latin-1 but not valid UTF-8
        latin1_content = b"\xe9\xe8\xe7"  # éèç in Latin-1

        text = await text_extractor.extract_text(
            file_data=latin1_content, file_type="text/plain"
        )

        assert isinstance(text, str)
        assert len(text) > 0
        print(f"Latin-1 text: {text}")

    @pytest.mark.asyncio
    async def test_large_file_handling(self, text_extractor):
        """Test handling of large text files."""
        large_content = ("This is a test line.\n" * 10000).encode("utf-8")

        text = await text_extractor.extract_text(
            file_data=large_content, file_type="text/plain"
        )

        assert isinstance(text, str)
        assert len(text) > 0
        assert "test line" in text
        print(f"Large file text length: {len(text)}")
