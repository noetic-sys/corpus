from __future__ import annotations

import asyncio
from typing import List
from io import BytesIO
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

from .interface import DocumentHighlightingInterface, HighlightData
from common.core.otel_axiom_exporter import trace_span, get_logger

logger = get_logger(__name__)


class WordHighlightProvider(DocumentHighlightingInterface):
    """Word-specific provider for highlighting text using python-docx."""

    def __init__(self):
        self.supported_extensions = [".docx", ".doc"]

    @trace_span
    async def highlight_document(
        self,
        document_content: bytes,
        document_filename: str,
        highlights: List[HighlightData],
    ) -> bytes:
        """Highlight text in a Word document using python-docx."""
        if not self.supports_file_type(document_filename):
            raise NotImplementedError(f"File type not supported: {document_filename}")

        if not highlights:
            return document_content

        # Run the synchronous Word processing in a thread pool
        return await asyncio.to_thread(
            self._highlight_word_sync, document_content, highlights
        )

    def _highlight_word_sync(
        self, document_content: bytes, highlights: List[HighlightData]
    ) -> bytes:
        """Synchronous Word highlighting logic."""
        try:
            # Open Word document from bytes
            input_stream = BytesIO(document_content)
            doc = Document(input_stream)
            input_stream.close()

            # Track processed highlights to avoid duplicates
            processed_texts = set()
            total_highlights_applied = 0

            for highlight_data in highlights:
                text_to_find = highlight_data.text.strip()

                # Skip if we've already processed this exact text
                if text_to_find in processed_texts:
                    continue
                processed_texts.add(text_to_find)

                logger.info(
                    f"Searching for text to highlight: '{text_to_find[:50]}...'"
                )

                # Search for text across all paragraphs
                found_instances = 0
                for paragraph in doc.paragraphs:
                    found_instances += self._highlight_text_in_paragraph(
                        paragraph, text_to_find, highlight_data
                    )

                # Search for text in tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                found_instances += self._highlight_text_in_paragraph(
                                    paragraph, text_to_find, highlight_data
                                )

                # Search for text in headers and footers
                for section in doc.sections:
                    # Headers
                    if section.header:
                        for paragraph in section.header.paragraphs:
                            found_instances += self._highlight_text_in_paragraph(
                                paragraph, text_to_find, highlight_data
                            )

                    # Footers
                    if section.footer:
                        for paragraph in section.footer.paragraphs:
                            found_instances += self._highlight_text_in_paragraph(
                                paragraph, text_to_find, highlight_data
                            )

                if found_instances == 0:
                    logger.warning(
                        f"Text not found in document: '{text_to_find[:50]}...'"
                    )
                else:
                    logger.info(
                        f"Highlighted {found_instances} instances of text: '{text_to_find[:50]}...'"
                    )
                    total_highlights_applied += found_instances

            # Save the modified Word document to bytes
            output_stream = BytesIO()
            doc.save(output_stream)
            highlighted_content = output_stream.getvalue()
            output_stream.close()

            logger.info(
                f"Successfully highlighted Word document with {len(highlights)} highlight requests, {total_highlights_applied} instances highlighted"
            )
            return highlighted_content

        except Exception as e:
            logger.error(f"Error highlighting Word document: {e}")
            raise

    def _highlight_text_in_paragraph(
        self, paragraph, text_to_find: str, highlight_data: HighlightData
    ) -> int:
        """
        Highlight text within a paragraph.

        Returns:
            Number of instances highlighted
        """
        found_instances = 0
        paragraph_text = paragraph.text

        if text_to_find.lower() not in paragraph_text.lower():
            return 0

        # Find all occurrences in this paragraph
        start_pos = 0
        while True:
            pos = paragraph_text.lower().find(text_to_find.lower(), start_pos)
            if pos == -1:
                break

            # Try to highlight the text by modifying runs
            success = self._highlight_text_in_runs(
                paragraph, pos, pos + len(text_to_find), highlight_data
            )
            if success:
                found_instances += 1
            start_pos = pos + 1

        return found_instances

    def _highlight_text_in_runs(
        self, paragraph, start_char: int, end_char: int, highlight_data: HighlightData
    ) -> bool:
        """
        Highlight specific character range in paragraph runs.

        Returns:
            True if highlighting was successfully applied
        """
        try:
            # Convert RGB (0-1) to RGB (0-255) for Word
            rgb_color = RGBColor(
                int(highlight_data.color[0] * 255),
                int(highlight_data.color[1] * 255),
                int(highlight_data.color[2] * 255),
            )

            # Find the runs that contain our target text
            current_pos = 0
            for run in paragraph.runs:
                run_length = len(run.text)
                run_start = current_pos
                run_end = current_pos + run_length

                # Check if this run overlaps with our target range
                if run_start < end_char and run_end > start_char:
                    # Apply highlighting
                    try:
                        # Try to set highlight color (yellow is most common/supported)
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    except Exception:
                        try:
                            # Fallback: try to set background color if supported
                            run.font.color.rgb = rgb_color
                        except Exception:
                            # Last fallback: just change text color
                            logger.warning(
                                "Using text color change as highlighting fallback"
                            )
                            run.font.color.rgb = rgb_color

                current_pos = run_end

            return True

        except Exception as e:
            logger.warning(f"Could not apply highlighting to paragraph: {e}")
            return False

    def supports_file_type(self, filename: str) -> bool:
        """Check if the file is a Word document."""
        return any(filename.lower().endswith(ext) for ext in self.supported_extensions)

    def get_supported_extensions(self) -> List[str]:
        """Return list of supported file extensions."""
        return self.supported_extensions.copy()
