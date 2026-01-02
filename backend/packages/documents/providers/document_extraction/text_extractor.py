import io
from typing import Dict, Any
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

from .interface import DocumentExtractorInterface
from common.core.otel_axiom_exporter import get_logger
from common.providers.storage.factory import get_storage
from packages.documents.models.domain.document import DocumentModel

logger = get_logger(__name__)


class TextExtractor(DocumentExtractorInterface):
    """Document text extractor supporting PDF, DOCX, and TXT files."""

    SUPPORTED_TYPES = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "text/plain": "txt",
        "text/markdown": "md",
    }

    def __init__(self):
        self.storage = get_storage()

    def supports_file_type(self, file_type: str) -> bool:
        """Check if the extractor supports the given file type."""
        return file_type.lower() in self.SUPPORTED_TYPES

    async def extract_text(self, document: DocumentModel) -> str:
        """Extract text content from a document file."""
        try:
            # Download file from storage
            file_data = await self.storage.download(document.storage_key)
            if not file_data:
                raise ValueError(f"Failed to download file from {document.storage_key}")

            file_type = document.content_type or ""
            file_type_lower = file_type.lower()

            if file_type_lower == "application/pdf":
                return await self._extract_pdf_text(file_data)
            elif (
                file_type_lower
                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                return await self._extract_docx_text(file_data)
            elif file_type_lower in ["text/plain", "text/markdown"]:
                return await self._extract_plain_text(file_data)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")

        except Exception as e:
            logger.error(f"Error extracting text from {file_type}: {e}")
            raise Exception(f"Failed to extract text: {str(e)}")

    async def get_metadata(self, document: DocumentModel) -> Dict[str, Any]:
        """Extract metadata from a document file."""
        try:
            # Download file from storage
            file_data = await self.storage.download(document.storage_key)
            if not file_data:
                raise ValueError(f"Failed to download file from {document.storage_key}")

            file_type = document.content_type or ""
            metadata = {
                "file_size": len(file_data),
                "file_type": file_type,
            }

            file_type_lower = file_type.lower()

            if file_type_lower == "application/pdf":
                pdf_metadata = await self._extract_pdf_metadata(file_data)
                metadata.update(pdf_metadata)
            elif (
                file_type_lower
                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                docx_metadata = await self._extract_docx_metadata(file_data)
                metadata.update(docx_metadata)

            return metadata

        except Exception as e:
            logger.error(f"Error extracting metadata from {file_type}: {e}")
            return {
                "file_size": document.file_size or 0,
                "file_type": document.content_type or "",
                "error": str(e),
            }

    async def _extract_pdf_text(self, file_data: bytes) -> str:
        """Extract text from PDF file."""
        pdf_file = io.BytesIO(file_data)
        reader = PdfReader(pdf_file)

        text_parts = []
        for page in reader.pages:
            text_parts.append(page.extract_text())

        full_text = "\n\n".join(text_parts)
        logger.info(f"Extracted text from PDF ({len(reader.pages)} pages)")
        return full_text.strip()

    async def _extract_docx_text(self, file_data: bytes) -> str:
        """Extract text from DOCX file."""
        docx_file = io.BytesIO(file_data)
        doc = DocxDocument(docx_file)

        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        full_text = "\n\n".join(text_parts)
        logger.info(f"Extracted text from DOCX ({len(doc.paragraphs)} paragraphs)")
        return full_text.strip()

    async def _extract_plain_text(self, file_data: bytes) -> str:
        """Extract text from plain text file."""
        try:
            text = file_data.decode("utf-8")
        except UnicodeDecodeError:
            # Fallback to latin-1 if UTF-8 fails
            text = file_data.decode("latin-1")

        logger.info("Extracted plain text content")
        return text.strip()

    async def _extract_pdf_metadata(self, file_data: bytes) -> Dict[str, Any]:
        """Extract metadata from PDF file."""
        pdf_file = io.BytesIO(file_data)
        reader = PdfReader(pdf_file)

        metadata = {}
        if reader.metadata:
            metadata.update(
                {
                    "title": reader.metadata.get("/Title", ""),
                    "author": reader.metadata.get("/Author", ""),
                    "subject": reader.metadata.get("/Subject", ""),
                    "creator": reader.metadata.get("/Creator", ""),
                    "producer": reader.metadata.get("/Producer", ""),
                    "creation_date": str(reader.metadata.get("/CreationDate", "")),
                    "modification_date": str(reader.metadata.get("/ModDate", "")),
                }
            )

        metadata["page_count"] = len(reader.pages)
        return metadata

    async def _extract_docx_metadata(self, file_data: bytes) -> Dict[str, Any]:
        """Extract metadata from DOCX file."""
        docx_file = io.BytesIO(file_data)
        doc = DocxDocument(docx_file)

        metadata = {}
        if doc.core_properties:
            props = doc.core_properties
            metadata.update(
                {
                    "title": props.title or "",
                    "author": props.author or "",
                    "subject": props.subject or "",
                    "keywords": props.keywords or "",
                    "category": props.category or "",
                    "comments": props.comments or "",
                    "created": str(props.created) if props.created else "",
                    "modified": str(props.modified) if props.modified else "",
                    "last_modified_by": props.last_modified_by or "",
                }
            )

        metadata["paragraph_count"] = len(doc.paragraphs)
        return metadata
